# ==========================================
# PROTOTIPO DE AUTODIAGNÓSTICO ENS
# Trabajo Fin de Estudios - Ingeniería Informática
# Autora: Mónica Polixenia Palasanu
# ==========================================

from docx import Document
from datetime import datetime


# -----------------------------
# Conversión de valores
# -----------------------------

def valor_a_numero(valor):
    valores = {
        "BAJO": 1,
        "MEDIO": 2,
        "ALTO": 3
    }
    return valores.get(valor.upper(), 0)


def numero_a_categoria(valor):
    if valor == 3:
        return "ALTA"
    elif valor == 2:
        return "MEDIA"
    elif valor == 1:
        return "BASICA"
    return "NO DETERMINADA"


# -----------------------------
# Categorización ENS
# -----------------------------

def calcular_categoria_ens(
        disponibilidad,
        confidencialidad,
        integridad,
        autenticidad,
        trazabilidad):

    valores = [
        valor_a_numero(disponibilidad),
        valor_a_numero(confidencialidad),
        valor_a_numero(integridad),
        valor_a_numero(autenticidad),
        valor_a_numero(trazabilidad)
    ]

    valor_maximo = max(valores)

    return numero_a_categoria(valor_maximo)


# -----------------------------
# Medidas aplicables
# -----------------------------

def obtener_medidas_aplicables(categoria):

    medidas_basicas = [
        "ORG.1 - Política de seguridad",
        "ORG.2 - Normativa de seguridad",
        "ORG.3 - Procedimientos de seguridad",
        "OP.EXP.1 - Inventario de activos",
        "OP.EXP.7 - Gestión de incidentes"
    ]

    medidas_medias = [
        "OP.CONT.1 - Análisis de impacto",
        "OP.CONT.2 - Plan de continuidad",
        "MP.INFO.2 - Calificación de la información",
        "MP.SI.2 - Protección frente a código dañino"
    ]

    medidas_altas = [
        "OP.MON.1 - Detección de intrusión",
        "OP.MON.2 - Sistema de métricas",
        "OP.MON.3 - Vigilancia",
        "MP.COM.3 - Protección de comunicaciones"
    ]

    if categoria == "BASICA":
        return medidas_basicas

    elif categoria == "MEDIA":
        return medidas_basicas + medidas_medias

    elif categoria == "ALTA":
        return medidas_basicas + medidas_medias + medidas_altas

    return []


# -----------------------------
# Cumplimiento inicial
# -----------------------------

def calcular_cumplimiento(respuestas):

    total = len(respuestas)
    cumplidas = respuestas.count("SI")

    if total == 0:
        return 0

    return round((cumplidas / total) * 100, 2)


# -----------------------------
# Generación Word
# -----------------------------

def generar_documento(
        nombre_empresa,
        categoria,
        cumplimiento,
        medidas):

    documento = Document()

    documento.add_heading(
        'INFORME PRELIMINAR DE AUTODIAGNÓSTICO ENS',
        level=1
    )

    documento.add_paragraph(
        f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y %H:%M")}'
    )

    documento.add_paragraph(
        f'Empresa evaluada: {nombre_empresa}'
    )

    documento.add_paragraph(
        f'Categoría ENS preliminar: {categoria}'
    )

    documento.add_paragraph(
        f'Nivel inicial de cumplimiento: {cumplimiento}%'
    )

    documento.add_heading(
        'Medidas aplicables identificadas',
        level=2
    )

    for medida in medidas:
        documento.add_paragraph(
            medida,
            style='List Bullet'
        )

    documento.add_heading(
        'Conclusión preliminar',
        level=2
    )

    documento.add_paragraph(
        'El resultado obtenido permite identificar de forma inicial '
        'la categoría ENS aplicable, las medidas de seguridad que '
        'deberían ser consideradas y el grado aproximado de '
        'cumplimiento inicial. Este resultado debe ser revisado '
        'posteriormente por personal experto antes de considerarse '
        'definitivo.'
    )

    documento.save(
        "informe_autodiagnostico_ens.docx"
    )


# -----------------------------
# Programa principal
# -----------------------------

def main():

    print("\n====================================")
    print(" PROTOTIPO DE AUTODIAGNÓSTICO ENS")
    print("====================================\n")

    nombre_empresa = input(
        "Nombre de la empresa: "
    )

    disponibilidad = input(
        "Impacto en disponibilidad (BAJO/MEDIO/ALTO): "
    )

    confidencialidad = input(
        "Impacto en confidencialidad (BAJO/MEDIO/ALTO): "
    )

    integridad = input(
        "Impacto en integridad (BAJO/MEDIO/ALTO): "
    )

    autenticidad = input(
        "Impacto en autenticidad (BAJO/MEDIO/ALTO): "
    )

    trazabilidad = input(
        "Impacto en trazabilidad (BAJO/MEDIO/ALTO): "
    )

    categoria = calcular_categoria_ens(
        disponibilidad,
        confidencialidad,
        integridad,
        autenticidad,
        trazabilidad
    )

    medidas = obtener_medidas_aplicables(
        categoria
    )

    print("\nResponda SI o NO:\n")

    respuestas = []

    respuestas.append(
        input("¿Existe política de seguridad aprobada? ")
    )

    respuestas.append(
        input("¿Existe inventario de activos? ")
    )

    respuestas.append(
        input("¿Existe análisis de riesgos? ")
    )

    respuestas.append(
        input("¿Existe gestión de incidencias? ")
    )

    respuestas.append(
        input("¿Existe plan de continuidad? ")
    )

    respuestas = [
        r.upper() for r in respuestas
    ]

    cumplimiento = calcular_cumplimiento(
        respuestas
    )

    generar_documento(
        nombre_empresa,
        categoria,
        cumplimiento,
        medidas
    )

    print("\n====================================")
    print(" RESULTADOS")
    print("====================================\n")

    print(f"Empresa: {nombre_empresa}")
    print(f"Categoría ENS: {categoria}")
    print(f"Nivel de cumplimiento: {cumplimiento}%")

    print("\nMedidas aplicables:")

    for medida in medidas:
        print(f"- {medida}")

    print(
        "\nInforme generado correctamente:"
    )

    print(
        "informe_autodiagnostico_ens.docx"
    )


if __name__ == "__main__":
    main()
